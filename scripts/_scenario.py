"""Shared setup for the manual test scripts.

Importing this module points the app at a throwaway database, so it MUST be
imported before anything from ``app.db`` — ``app.db.session`` reads the path at
import time. Set PLATER_DB yourself beforehand to run against a database you
want to keep and inspect.
"""

from __future__ import annotations

import os
import sys
import tempfile
from datetime import date
from decimal import Decimal
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

SCRATCH = Path(tempfile.mkdtemp(prefix="plater_script_"))
os.environ.setdefault("PLATER_DB", str(SCRATCH / "plater.db"))

OUTPUT_DIR = PROJECT_ROOT / "scripts" / "output"

# seeded reference codes these scenarios rely on
TAX_SYSTEM = "ua_edrpou"
UNIT = "hour"
CURRENCY = "UAH"
COUNTRY = "UKR"
DOCUMENT_TYPE = "invoice"


def banner(text: str) -> None:
    print(f"\n{'=' * 78}\n{text}\n{'=' * 78}")


def initialise() -> None:
    """Migrate and seed the throwaway database (reference data + shipped templates)."""

    from app.db.session import engine, init_db

    engine.echo = False
    init_db()

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    print(f"database: {os.environ['PLATER_DB']}")


def default_template_config(
    *,
    primary: str = "ENG",
    secondary: str | None = "UKR",
    document_type: str = DOCUMENT_TYPE,
    name: str = "Manual test",
):
    """DefaultTemplateConfig has no seed — it is user configuration, not reference data."""

    from app.db.models.configs.default_template_config import DefaultTemplateConfig

    return DefaultTemplateConfig(
        primary_language_code=primary,
        secondary_language_code=secondary,
        document_type_code=document_type,
        name=name,
        description="created by scripts/_scenario.py",
        append_currency=True,
    )


def sample_organisation(session, en_name: str, uk_name: str, tax_value: str,
                        *, with_bank: bool = True):
    """A provider/client with localizations, a tax id, a representative and a bank."""

    from app.db.models.core.bank_account import BankAccount
    from app.db.models.core.bank_account_localization import BankAccountLocalization
    from app.db.models.core.organization import Organization
    from app.db.models.core.organization_localization import OrganizationLocalization
    from app.db.models.core.representative import Representative
    from app.db.models.core.representative_localization import RepresentativeLocalization
    from app.db.models.core.tax_id import TaxId

    org = Organization(
        email=f"{en_name.split()[0].lower()}@example.com",
        phone="+380441234567",
        localizations={
            "ENG": OrganizationLocalization(
                language_code="ENG", org_type="LLC",
                legal_name=en_name, address="1 Main St, Kyiv",
            ),
            "UKR": OrganizationLocalization(
                language_code="UKR", org_type="ТОВ",
                legal_name=uk_name, address="вул. Головна 1, Київ",
            ),
        },
    )

    org.tax_ids.append(
        TaxId(tax_id_system_code=TAX_SYSTEM, country_code=COUNTRY, value=tax_value)
    )
    org.representatives.append(Representative(localizations={
        "ENG": RepresentativeLocalization(
            language_code="ENG", name="Ivan Petrenko", title="Director",
        ),
        "UKR": RepresentativeLocalization(
            language_code="UKR", name="Іван Петренко", title="Директор",
        ),
    }))

    if with_bank:
        org.bank_accounts.append(BankAccount(
            country_code=COUNTRY, currency_code=CURRENCY,
            iban="UA903052990000026007233566001", swift="PBANUA2X",
            localizations={
                "ENG": BankAccountLocalization(
                    language_code="ENG", bank_name="PrivatBank", bank_info="MFO 305299",
                ),
                "UKR": BankAccountLocalization(
                    language_code="UKR", bank_name="ПриватБанк", bank_info="МФО 305299",
                ),
            },
        ))

    session.add(org)
    return org


# def sample_invoice_lines(session) -> list:
#     """Three lines with tax, enough to exercise totals and row expansion."""

#     from app.db.models.core.invoice_line import InvoiceLine
#     from app.db.models.core.invoice_line_localization import InvoiceLineLocalization

#     rows = [
#         ("Backend development", "Розробка бекенду", "40.000", "85.00", "0.20000"),
#         ("UI/UX design", "UI/UX дизайн", "12.500", "95.50", "0.20000"),
#         ("Consulting", "Консультації", "3.000", "150.00", "0.20000"),
#     ]

#     lines = []
#     for desc_en, desc_uk, quantity, price, tax in rows:
#         line = InvoiceLine(
#             quantity=Decimal(quantity),
#             unit_price=Decimal(price),
#             tax_rate=Decimal(tax),
#             measurement_unit_code=UNIT,
#             localizations={
#                 "ENG": InvoiceLineLocalization(language_code="ENG", description=desc_en),
#                 "UKR": InvoiceLineLocalization(language_code="UKR", description=desc_uk),
#             },
#         )
#         session.add(line)
#         lines.append(line)

#     return lines


def sample_lines() -> tuple:
    """Three lines with tax, enough to exercise totals and row expansion.
    Values now, not stored rows — the draft carries them directly.
    """

    from app.services.invoice.draft import LineInput

    rows = [
        ("Backend development", "Розробка бекенду", "40.000", "85.00", "0.20000"),
        ("UI/UX design", "UI/UX дизайн", "12.500", "95.50", "0.20000"),
        ("Consulting", "Консультації", "3.000", "150.00", "0.20000"),
    ]

    return tuple(
        LineInput(
            descriptions={"ENG": desc_en, "UKR": desc_uk},
            unit_code=UNIT,
            quantity=Decimal(quantity),
            unit_price=Decimal(price),
            tax_rate=Decimal(tax),
        )
        for desc_en, desc_uk, quantity, price, tax in rows
    )


def sample_draft(provider, client, sequence, issue_date: date | None = None):
    """An InvoiceDraft selecting the provider's representative and bank account."""

    from app.services.invoice.draft import InvoiceDraft, PartySelection

    return InvoiceDraft(
        template_id=0,                       # not used by the assembler
        sequence_id=sequence.id,
        currency_code=CURRENCY,
        issue_date=issue_date or date.today(),
        provider=PartySelection(
            organization_id=provider.id,
            tax_id_id=provider.tax_ids[0].id,
            representative_id=provider.representatives[0].id,
            bank_account_id=provider.bank_accounts[0].id,
        ),
        client=PartySelection(
            organization_id=client.id,
            tax_id_id=client.tax_ids[0].id,
        ),
        lines=sample_lines(),
    )


def describe_docx(path: Path) -> str:
    """Every piece of text in a rendered document — body, tables, headers, footers."""

    from docx import Document

    document = Document(str(path))
    parts = [p.text for p in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)

    for section in document.sections:
        for area in (section.header, section.footer):
            parts.extend(p.text for p in area.paragraphs)
            for table in area.tables:
                for row in table.rows:
                    parts.extend(cell.text for cell in row.cells)

    return "\n".join(parts)


def print_document(path: Path) -> None:
    from docx import Document

    document = Document(str(path))
    print(f"   paragraphs={len(document.paragraphs)} tables={len(document.tables)} "
          f"images={len(document.inline_shapes)}")

    for index, table in enumerate(document.tables):
        print(f"   table {index}: {len(table.rows)}x{len(table.columns)}")
        for row in table.rows[:6]:
            cells = " | ".join(c.text.replace("\n", " ")[:28] for c in row.cells)
            print(f"      | {cells}")

    for index, section in enumerate(document.sections):
        header = "\n".join(p.text for p in section.header.paragraphs).strip()
        footer = "\n".join(p.text for p in section.footer.paragraphs).strip()
        if header:
            print(f"   section {index} header: {header!r}")
        if footer:
            print(f"   section {index} footer: {footer!r}")

    print("\n   --- body ---")
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            print(f"   {paragraph.text}")
