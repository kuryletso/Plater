"""A TABLE placeholder written on its own line becomes a table of its own."""

import pytest
from docx import Document

from app.core.diagnostics import DiagnosticCollector
from app.document_engine.blueprint.models.paragraph import ParagraphBlueprint
from app.document_engine.blueprint.models.table import TableBlueprint, TablePlaceholder
from app.document_engine.enums.enums import PlaceholderType, ResolveMode
from app.document_engine.orchestration.pipeline import (
    TemplateIngestionPipeline, TemplateRenderingPipeline,
)
from app.document_engine.rendering.context import (
    InvoiceLineRow, InvoiceTableData, RenderContext,
)
from app.document_engine.rendering.resolve.models import ResolvedParagraph, ResolvedTable
from app.document_engine.rendering.resolve.resolver import DocumentResolver

from tests.conftest import FixtureInputProvider


class NoAssets:
    def get(self, asset_id):
        return None


def table_provider():
    return FixtureInputProvider(placeholders={
        "org_name": {"active": True, "required": True, "type": PlaceholderType.SCALAR},
        "invoice_table": {"active": True, "required": False, "type": PlaceholderType.TABLE},
    })


def ingest(path, provider=None):
    pipeline = TemplateIngestionPipeline(provider or table_provider())
    result = pipeline.ingest(path)
    return pipeline.finalize(result.draft), result.diagnostics


def blocks_of(blueprint) -> list:
    return [block for section in blueprint.sections for block in section.blocks]


def sample_table_data(language: str = "ENG") -> InvoiceTableData:
    columns = ("invl_n", "invl_desc", "invl_unit", "invl_qnty", "invl_price", "invl_total")
    values = {column: {language: f"<{column}>"} for column in columns}

    return InvoiceTableData(
        rows=(InvoiceLineRow(values=values),),
        show_tax=False,
        subtotal={language: "100.00"},
        total_tax=None,
        total={language: "100.00"},
        labels={key: {language: key} for key in (*columns, "subtotal", "total_tax", "total")},
    )


# --- promotion ---------------------------------------------------------------

def test_a_lone_table_placeholder_becomes_a_table_block(make_docx):
    blueprint, diagnostics = ingest(make_docx(paragraphs=["{{ invoice_table }}"]))

    assert [type(b) for b in blocks_of(blueprint)] == [TablePlaceholder]
    assert diagnostics.warnings == []


def test_the_promoted_block_keeps_its_key_and_language(make_docx):
    blueprint, _ = ingest(make_docx(paragraphs=["{{ invoice_table.UKR }}"]))

    block = blocks_of(blueprint)[0]
    assert block.key == "invoice_table"
    assert block.language == "UKR"


def test_a_standalone_table_gets_visible_default_borders(make_docx):
    """A borderless invoice table reads as broken, so the fallback draws them."""
    blueprint, _ = ingest(make_docx(paragraphs=["{{ invoice_table }}"]))

    style = blocks_of(blueprint)[0].style
    assert style.border_top.size > 0
    assert style.border_inside_h.size > 0


def test_surrounding_whitespace_does_not_prevent_promotion(make_docx):
    blueprint, diagnostics = ingest(make_docx(paragraphs=["   {{ invoice_table }}  "]))

    assert [type(b) for b in blocks_of(blueprint)] == [TablePlaceholder]
    assert diagnostics.warnings == []


def test_a_table_placeholder_beside_text_warns_instead(make_docx):
    """It cannot become a table mid-sentence; say so at ingestion, not at render."""
    blueprint, diagnostics = ingest(make_docx(paragraphs=["Total: {{ invoice_table }}"]))

    assert [type(b) for b in blocks_of(blueprint)] == [ParagraphBlueprint]
    assert [w.code for w in diagnostics.warnings] == ["table_placeholder_not_alone"]


def test_an_ordinary_paragraph_is_untouched(make_docx):
    blueprint, _ = ingest(make_docx(paragraphs=["Invoice for {{ org_name }}"]))

    assert [type(b) for b in blocks_of(blueprint)] == [ParagraphBlueprint]


def test_a_wrapped_table_placeholder_still_inherits_the_wrapper_style(make_docx):
    """The in-table path must keep working — that is how a template supplies a style."""
    blueprint, _ = ingest(make_docx(table=[["{{ invoice_table }}"]]))

    blocks = blocks_of(blueprint)
    assert [type(b) for b in blocks] == [TablePlaceholder]
    assert not isinstance(blocks[0], TableBlueprint)


# --- rendering ---------------------------------------------------------------

def test_a_standalone_placeholder_renders_the_generated_table(make_docx, tmp_path):
    blueprint, _ = ingest(make_docx(paragraphs=["{{ invoice_table }}"]))
    context = RenderContext(scalars={}, table=sample_table_data())

    render = TemplateRenderingPipeline(NoAssets()).render(blueprint, context)

    assert render.docx is not None
    out = tmp_path / "standalone.docx"
    out.write_bytes(render.docx)

    document = Document(str(out))
    assert len(document.tables) == 1
    assert len(document.tables[0].rows) >= 2          # header + a line + totals


def test_the_generated_table_is_not_left_as_an_empty_scalar(make_docx):
    """Regression: it used to fall through to a scalar lookup and render empty."""
    blueprint, _ = ingest(make_docx(paragraphs=["{{ invoice_table }}"]))
    context = RenderContext(scalars={}, table=sample_table_data())

    render = TemplateRenderingPipeline(NoAssets()).render(blueprint, context)

    assert [d.code for d in render.diagnostics.items] == []


# --- raw rendering -----------------------------------------------------------

def test_raw_rendering_gives_back_the_placeholder_not_a_table(make_docx):
    """The author wrote one placeholder, so a raw render must return one placeholder."""
    blueprint, _ = ingest(make_docx(paragraphs=["{{ invoice_table }}"]))

    resolved = DocumentResolver(
        RenderContext(), NoAssets(), DiagnosticCollector(), ResolveMode.KEYS,
    ).resolve(blueprint)

    block = resolved.sections[0].blocks[0]
    assert isinstance(block, ResolvedParagraph)
    assert not isinstance(block, ResolvedTable)
    assert block.runs[0].text == "{{ invoice_table }}"
    assert block.runs[0].placeholder_key == "invoice_table"


def test_a_wrapped_placeholder_also_raw_renders_as_text(make_docx):
    blueprint, _ = ingest(make_docx(table=[["{{ invoice_table }}"]]))

    resolved = DocumentResolver(
        RenderContext(), NoAssets(), DiagnosticCollector(), ResolveMode.KEYS,
    ).resolve(blueprint)

    assert resolved.sections[0].blocks[0].runs[0].text == "{{ invoice_table }}"


def test_a_raw_rendered_table_placeholder_can_be_re_ingested(make_docx, tmp_path):
    """The payoff: raw export round-trips back into a working table placeholder."""
    blueprint, _ = ingest(make_docx(paragraphs=["{{ invoice_table }}"]))

    render = TemplateRenderingPipeline(NoAssets()).render_raw(blueprint)
    exported = tmp_path / "exported.docx"
    exported.write_bytes(render.docx)

    reimported, diagnostics = ingest(exported)

    assert [type(b) for b in blocks_of(reimported)] == [TablePlaceholder]
    assert diagnostics.warnings == []
