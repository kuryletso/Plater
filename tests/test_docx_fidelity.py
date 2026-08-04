"""Fidelity of real-world Word documents through parse -> render.

Every case here was found by ingesting an actual Word template; python-docx writes
much cleaner OOXML than Word does, so none of these are reachable from synthetic
documents alone.
"""

from lxml import etree

from app.core.diagnostics import DiagnosticCollector
from app.document_engine.blueprint.models.segment import TextStyleBlueprint
from app.document_engine.parser.models.blocks import ParagraphNode
from app.document_engine.parser.models.inlines import RunNode
from app.document_engine.parser.parser import DocxParser
from app.document_engine.rendering.docx.run import build_run
from app.document_engine.rendering.docx.xml import qn


def parse(path) -> list:
    with DocxParser(path, diagnostics=DiagnosticCollector()) as parser:
        return parser.parse()


def runs_of(blocks) -> list[RunNode]:
    return [
        node
        for block in blocks
        if isinstance(block, ParagraphNode)
        for node in block.inlines
        if isinstance(node, RunNode)
    ]


def style(**overrides) -> TextStyleBlueprint:
    base = dict(bold=False, italic=False, underline=False,
                font_name="Calibri", font_size=20, color="000000")
    return TextStyleBlueprint(**(base | overrides))


def children_tags(run: etree._Element) -> list[str]:
    return [etree.QName(child).localname for child in run if etree.QName(child).localname != "rPr"]


def table_style(**overrides):
    from app.document_engine.blueprint.models.margins import MarginsBlueprint
    from app.document_engine.blueprint.models.table import (
        TableBorderBlueprint, TableStyleBlueprint, TableWidthBlueprint,
    )
    from app.document_engine.enums.enums import TableBorderStyleEnum, TableWidthType

    border = TableBorderBlueprint(
        style=TableBorderStyleEnum.SINGLE, size=4, color="000000",
    )
    base = dict(
        width=TableWidthBlueprint(value=9600, type=TableWidthType.DXA),
        autofit=False,
        border_top=border, border_bottom=border,
        border_left=border, border_right=border,
        border_inside_v=border, border_inside_h=border,
        margins=MarginsBlueprint(top=0, bottom=0, left=100, right=100),
    )
    return TableStyleBlueprint(**(base | overrides))


# --- manual line breaks: parsing ---------------------------------------------

def test_manual_break_is_parsed_as_newline(make_break_docx):
    """Shift+Enter emits <w:br/>, which must survive as '\\n' in the run text."""
    blocks = parse(make_break_docx(["line one", "line two"]))

    assert runs_of(blocks)[0].text == "line one\nline two"


def test_consecutive_breaks_are_preserved(make_break_docx):
    blocks = parse(make_break_docx(["a", "", "b"]))

    assert runs_of(blocks)[0].text == "a\n\nb"


def test_leading_break_is_preserved(make_break_docx):
    """A blank first line — the 'empty row' case."""
    blocks = parse(make_break_docx(["", "after blank"]))

    assert runs_of(blocks)[0].text == "\nafter blank"


# --- manual line breaks: emitting --------------------------------------------

def test_emitted_run_splits_newlines_into_break_elements():
    """OOXML renders '\\n' inside <w:t> as a space; a break must be a <w:br/> element."""
    run = build_run("line one\nline two", style())

    assert children_tags(run) == ["t", "br", "t"]
    assert [t.text for t in run.findall(qn("w:t"))] == ["line one", "line two"]


def test_emitted_run_without_newline_has_a_single_text_element():
    run = build_run("no breaks here", style())

    assert children_tags(run) == ["t"]


def test_emitted_blank_line_produces_a_lone_break():
    run = build_run("\nafter blank", style())

    assert children_tags(run) == ["br", "t"]


def test_emitted_consecutive_breaks_produce_consecutive_elements():
    run = build_run("a\n\nb", style())

    assert children_tags(run) == ["t", "br", "br", "t"]


def test_emitted_trailing_break_is_kept():
    run = build_run("text\n", style())

    assert children_tags(run) == ["t", "br"]


def test_break_round_trips_from_source_to_output(make_break_docx):
    """End to end: <w:br/> in the source survives parse and comes back as <w:br/>."""
    blocks = parse(make_break_docx(["first", "second"]))
    text = runs_of(blocks)[0].text

    assert children_tags(build_run(text, style())) == ["t", "br", "t"]


# --- font size inheritance ---------------------------------------------------

def test_run_inherits_size_from_the_default_paragraph_style(make_styled_docx):
    """Word puts the body size on Normal; docDefaults often carries none at all."""
    blocks = parse(make_styled_docx([("body text", None)], normal_half_points=20))

    assert runs_of(blocks)[0].style.font_size == 20


def test_direct_run_size_overrides_the_paragraph_style(make_styled_docx):
    blocks = parse(make_styled_docx([("big heading", 30)], normal_half_points=20))

    assert runs_of(blocks)[0].style.font_size == 30


def test_mixed_sizes_resolve_independently(make_styled_docx):
    """The real template's exact shape: mostly 10pt with a couple of 15pt lines."""
    blocks = parse(make_styled_docx(
        [("heading", 30), ("body", None), ("more body", None)],
        normal_half_points=20,
    ))

    assert [r.style.font_size for r in runs_of(blocks)] == [30, 20, 20]


def test_every_body_run_inherits_the_default_size(make_styled_docx):
    blocks = parse(make_styled_docx(
        [("one", None), ("two", None)], normal_half_points=20,
    ))

    assert [r.style.font_size for r in runs_of(blocks)] == [20, 20]


# --- cell borders ------------------------------------------------------------
#
# Layout tables routinely suppress borders per cell (w:val="nil"), so a cell border
# that is *absent* must stay absent — emitting a default would override the table's
# own borders on every cell.

def cell_borders_of(cell_element) -> dict[str, tuple[str, str, str]]:
    borders = cell_element.find(f"{qn('w:tcPr')}/{qn('w:tcBorders')}")
    if borders is None:
        return {}
    return {
        etree.QName(side).localname: (
            side.get(qn("w:val")), side.get(qn("w:sz")), side.get(qn("w:color")),
        )
        for side in borders
    }


def first_cell(table_element):
    return table_element.find(f"{qn('w:tr')}/{qn('w:tc')}")


def cell_style(**overrides):
    from app.document_engine.blueprint.models.margins import MarginsBlueprint
    from app.document_engine.blueprint.models.table import CellStyleBlueprint
    from app.document_engine.enums.enums import TableCellShading, VerticalAlignment

    base = dict(
        shading=TableCellShading.CLEAR,
        shading_fill="auto",
        margins=MarginsBlueprint(top=0, bottom=0, left=100, right=100),
        grid_span=1,
        v_alignment=VerticalAlignment.TOP,
    )
    return CellStyleBlueprint(**(base | overrides))


def border(style="single", size=8, color="000000"):
    from app.document_engine.blueprint.models.table import TableBorderBlueprint
    from app.document_engine.enums.enums import TableBorderStyleEnum

    return TableBorderBlueprint(
        style=TableBorderStyleEnum(style), size=size, color=color,
    )


def test_a_cell_without_borders_emits_no_tc_borders():
    """Absent means 'inherit from the table', so nothing may be written."""
    from app.document_engine.rendering.docx.table import build_cell

    cell = build_cell(blocks=[], style=cell_style())

    assert cell_borders_of(cell) == {}


def test_cell_borders_are_emitted_with_width_and_colour():
    from app.document_engine.rendering.docx.table import build_cell

    cell = build_cell(blocks=[], style=cell_style(
        border_top=border("single", 12, "FF0000"),
        border_bottom=border("single", 12, "FF0000"),
    ))

    assert cell_borders_of(cell) == {
        "top": ("single", "12", "FF0000"),
        "bottom": ("single", "12", "FF0000"),
    }


def test_a_suppressed_cell_border_round_trips_as_nil():
    """A layout cell hiding an inherited border must keep val='nil', not become 'none'."""
    from app.document_engine.rendering.docx.table import build_cell

    cell = build_cell(blocks=[], style=cell_style(
        border_top=border("nil", 0, "000000"),
        border_left=border("nil", 0, "000000"),
        border_bottom=border("nil", 0, "000000"),
        border_right=border("nil", 0, "000000"),
    ))

    assert cell_borders_of(cell) == {
        side: ("nil", "0", "000000") for side in ("top", "left", "bottom", "right")
    }


def test_only_the_sides_that_were_set_are_emitted():
    from app.document_engine.rendering.docx.table import build_cell

    cell = build_cell(blocks=[], style=cell_style(border_bottom=border()))

    assert list(cell_borders_of(cell)) == ["bottom"]


def test_tc_borders_precedes_shading_in_the_cell_properties():
    """CT_TcPr order is schema-significant: gridSpan, tcBorders, shd, tcMar, vAlign."""
    from app.document_engine.rendering.docx.table import build_cell

    cell = build_cell(blocks=[], style=cell_style(border_top=border()))
    children = [etree.QName(c).localname for c in cell.find(qn("w:tcPr"))]

    assert children.index("tcBorders") < children.index("shd")


def test_cell_borders_are_parsed_from_a_real_document(tmp_path):
    from docx import Document

    from app.document_engine.parser.models.blocks import TableNode

    document = Document()
    table = document.add_table(rows=1, cols=1)
    tc_pr = table.cell(0, 0)._tc.get_or_add_tcPr()
    borders = etree.SubElement(tc_pr, qn("w:tcBorders"))
    for side, val, size in (("w:top", "nil", "0"), ("w:bottom", "single", "18")):
        element = etree.SubElement(borders, qn(side))
        element.set(qn("w:val"), val)
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), "0000FF")

    path = tmp_path / "cell_borders.docx"
    document.save(path)

    node = next(b for b in parse(path) if isinstance(b, TableNode))
    style = node.rows[0].cells[0].style

    assert (style.border_top.style, style.border_top.size) == ("nil", 0)
    assert (style.border_bottom.style, style.border_bottom.size) == ("single", 18)
    assert style.border_bottom.color == "0000FF"
    assert style.border_left is None, "unset sides must stay unset"


def test_cell_borders_survive_parse_to_render(tmp_path):
    """End to end: a suppressed cell border must not reappear in the output."""
    from docx import Document

    from app.document_engine.normalization.structural_normalizer import StructuralNormalizer
    from app.document_engine.parser.models.blocks import TableNode

    document = Document()
    table = document.add_table(rows=1, cols=1)
    tc_pr = table.cell(0, 0)._tc.get_or_add_tcPr()
    borders = etree.SubElement(tc_pr, qn("w:tcBorders"))
    for side in ("w:top", "w:left", "w:bottom", "w:right"):
        element = etree.SubElement(borders, qn(side))
        element.set(qn("w:val"), "nil")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:color"), "000000")

    path = tmp_path / "layout.docx"
    document.save(path)

    sections = StructuralNormalizer.normalize(parse(path), DiagnosticCollector())
    normalized = next(b for b in sections[0].blocks if not hasattr(b, "inlines"))
    style = normalized.rows[0].cells[0].style

    assert style.border_top is not None
    assert style.border_top.style.value == "nil"


# --- table column widths -----------------------------------------------------

def grid_of(table_element) -> list[int]:
    grid = table_element.find(qn("w:tblGrid"))
    return [int(col.get(qn("w:w"))) for col in grid.findall(qn("w:gridCol"))]


def test_column_widths_are_parsed_from_the_table_grid(make_grid_docx):
    from app.document_engine.parser.models.blocks import TableNode

    blocks = parse(make_grid_docx([403, 3226, 1613]))
    table = next(b for b in blocks if isinstance(b, TableNode))

    assert table.style.column_widths == (403, 3226, 1613)


def test_partial_grid_is_ignored_rather_than_guessed(tmp_path, make_grid_docx):
    """A gridCol with no w attribute makes the whole grid untrustworthy."""
    import shutil
    import zipfile

    from app.document_engine.parser.models.blocks import TableNode

    source = make_grid_docx([403, 3226, 1613])
    stripped = tmp_path / "partial_grid.docx"

    with zipfile.ZipFile(source) as zin, zipfile.ZipFile(stripped, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                root = etree.fromstring(data)
                col = root.find(f".//{qn('w:tblGrid')}/{qn('w:gridCol')}")
                del col.attrib[qn("w:w")]
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8")
            zout.writestr(item, data)

    table = next(b for b in parse(stripped) if isinstance(b, TableNode))

    assert table.style.column_widths is None


def test_emitter_uses_the_source_widths_when_the_column_count_matches():
    from app.document_engine.rendering.docx.table import build_table

    style = table_style(column_widths=(403, 3226, 1613))
    table = build_table(rows=[], style=style, columns=3)

    assert grid_of(table) == [403, 3226, 1613]


def test_emitter_falls_back_to_even_split_on_column_count_mismatch():
    """The system-built invoice table sets its own column count; don't misapply a grid."""
    from app.document_engine.rendering.docx.table import build_table

    style = table_style(column_widths=(403, 3226, 1613))
    table = build_table(rows=[], style=style, columns=6)

    widths = grid_of(table)
    assert len(widths) == 6
    assert len(set(widths)) == 1, "mismatched grid must be ignored, not stretched"


def test_emitter_falls_back_to_even_split_without_a_grid():
    from app.document_engine.rendering.docx.table import build_table

    style = table_style(column_widths=())
    table = build_table(rows=[], style=style, columns=4)

    widths = grid_of(table)
    assert len(widths) == 4
    assert len(set(widths)) == 1


def test_uneven_widths_survive_parse_to_render(make_grid_docx):
    """End to end: the author's proportions must reach the emitted document."""
    from app.document_engine.parser.models.blocks import TableNode
    from app.document_engine.rendering.docx.table import build_table

    blocks = parse(make_grid_docx([403, 3226, 1613, 1613]))
    parsed = next(b for b in blocks if isinstance(b, TableNode))

    style = table_style(column_widths=parsed.style.column_widths)
    emitted = build_table(rows=[], style=style, columns=4)

    assert grid_of(emitted) == [403, 3226, 1613, 1613]
    assert grid_of(emitted) != [sum([403, 3226, 1613, 1613]) // 4] * 4


def test_default_style_size_applies_inside_table_cells(tmp_path):
    """Table-cell paragraphs resolve through the same path as body paragraphs."""
    from docx import Document
    from docx.shared import Pt

    document = Document()
    document.styles["Normal"].font.size = Pt(10)
    table = document.add_table(rows=1, cols=1)
    cell_paragraph = table.cell(0, 0).paragraphs[0]
    run = cell_paragraph.add_run("cell text")
    run.italic = False                     # force an rPr, as Word does

    path = tmp_path / "table_styled.docx"
    document.save(path)

    from app.document_engine.parser.models.blocks import TableNode

    cell_runs = [
        node
        for block in parse(path) if isinstance(block, TableNode)
        for row in block.rows
        for cell in row.cells
        for inner in cell.blocks if isinstance(inner, ParagraphNode)
        for node in inner.inlines if isinstance(node, RunNode)
    ]

    assert [r.style.font_size for r in cell_runs] == [20]
