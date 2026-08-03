"""Placeholder provenance on resolved runs, and KEYS-mode (raw) rendering."""

import pytest
from docx import Document

from app.core.diagnostics import DiagnosticCollector
from app.document_engine.blueprint.models.paragraph import ParagraphBlueprint
from app.document_engine.enums.enums import ResolveMode
from app.document_engine.orchestration.pipeline import (
    TemplateIngestionPipeline, TemplateRenderingPipeline,
)
from app.document_engine.rendering.context import RenderContext
from app.document_engine.rendering.resolve.models import ResolvedParagraph, ResolvedTextRun
from app.document_engine.rendering.resolve.raw import placeholder_syntax
from app.document_engine.rendering.resolve.resolver import DocumentResolver


class NoAssets:
    def get(self, asset_id):
        return None


def blueprint_of(path, provider):
    pipeline = TemplateIngestionPipeline(provider)
    return pipeline.finalize(pipeline.ingest(path).draft)


def resolve(blueprint, context=None, mode=ResolveMode.VALUES):
    return DocumentResolver(
        context or RenderContext(), NoAssets(), DiagnosticCollector(), mode,
    ).resolve(blueprint)


def runs_of(resolved) -> list[ResolvedTextRun]:
    return [
        run
        for section in resolved.sections
        for block in section.blocks
        if isinstance(block, ResolvedParagraph)
        for run in block.runs
        if isinstance(run, ResolvedTextRun)
    ]


def values(*keys) -> RenderContext:
    return RenderContext(
        scalars={key: {"ENG": f"<{key}>"} for key in keys},
    )


def column_provider():
    """A provider that knows invl_* as COLUMN, so rows become RowPlaceholders."""
    from tests.conftest import FixtureInputProvider
    from app.document_engine.enums.enums import PlaceholderType

    return FixtureInputProvider(placeholders={
        "org_name": {"active": True, "required": True, "type": PlaceholderType.SCALAR},
        "invoice_no": {"active": True, "required": True, "type": PlaceholderType.SCALAR},
        "invl_n": {"active": True, "required": False, "type": PlaceholderType.COLUMN},
        "invl_desc": {"active": True, "required": True, "type": PlaceholderType.COLUMN},
    })


def row_placeholders_in(blueprint) -> list:
    from app.document_engine.blueprint.models.table import RowPlaceholder, TableBlueprint

    return [
        row
        for section in blueprint.sections
        for block in section.blocks
        if isinstance(block, TableBlueprint)
        for row in block.rows
        if isinstance(row, RowPlaceholder)
    ]


# --- provenance --------------------------------------------------------------

def test_a_placeholder_run_carries_its_key(make_docx, fixture_provider):
    bp = blueprint_of(make_docx(paragraphs=["Hi {{ org_name }}"]), fixture_provider)

    runs = runs_of(resolve(bp, values("org_name")))

    assert [(r.text, r.placeholder_key) for r in runs] == [
        ("Hi ", None),
        ("<org_name>", "org_name"),
    ]


def test_literal_text_has_no_key(make_docx, fixture_provider):
    bp = blueprint_of(make_docx(paragraphs=["Just text"]), fixture_provider)

    assert [r.placeholder_key for r in runs_of(resolve(bp))] == [None]


def test_a_joined_placeholder_splits_into_one_run_per_key(make_docx, fixture_provider):
    """Per-placeholder highlighting needs each key in its own run."""
    bp = blueprint_of(
        make_docx(paragraphs=["{{ org_name, invoice_no, sep=', ' }}"]), fixture_provider,
    )

    runs = runs_of(resolve(bp, values("org_name", "invoice_no")))

    assert [(r.text, r.placeholder_key) for r in runs] == [
        ("<org_name>", "org_name"),
        (", ", None),
        ("<invoice_no>", "invoice_no"),
    ]


def test_a_joined_placeholder_still_drops_empty_parts(make_docx, fixture_provider):
    bp = blueprint_of(
        make_docx(paragraphs=["{{ org_name, invoice_no, sep=', ' }}"]), fixture_provider,
    )
    context = RenderContext(scalars={
        "org_name": {"ENG": "<org_name>"},
        "invoice_no": {"ENG": ""},
    })

    runs = runs_of(resolve(bp, context))

    assert [r.text for r in runs] == ["<org_name>"]


def test_a_grouped_placeholder_tags_each_member(make_docx, fixture_provider):
    bp = blueprint_of(
        make_docx(paragraphs=["{{ ((org_name) (invoice_no)) sep='; ' }}"]),
        fixture_provider,
    )

    runs = runs_of(resolve(bp, values("org_name", "invoice_no")))

    assert [(r.text, r.placeholder_key) for r in runs] == [
        ("<org_name>", "org_name"),
        ("; ", None),
        ("<invoice_no>", "invoice_no"),
    ]


def test_a_group_is_dropped_when_any_member_is_empty(make_docx, fixture_provider):
    """Existing semantics: an incomplete group prints nothing rather than half of itself."""
    bp = blueprint_of(
        make_docx(paragraphs=["{{ ((org_name invoice_no)) sep='; ' }}"]),
        fixture_provider,
    )
    context = RenderContext(scalars={
        "org_name": {"ENG": "<org_name>"},
        "invoice_no": {"ENG": ""},
    })

    assert [r.text for r in runs_of(resolve(bp, context))] == []


# --- KEYS mode ---------------------------------------------------------------

def test_keys_mode_emits_the_placeholder_syntax(make_docx, fixture_provider):
    bp = blueprint_of(make_docx(paragraphs=["Hi {{ org_name }}"]), fixture_provider)

    runs = runs_of(resolve(bp, mode=ResolveMode.KEYS))

    assert [r.text for r in runs] == ["Hi ", "{{ org_name }}"]


def test_keys_mode_needs_no_context_at_all(make_docx, fixture_provider):
    """A required placeholder would raise in VALUES mode with an empty context."""
    bp = blueprint_of(make_docx(paragraphs=["Hi {{ org_name }}"]), fixture_provider)

    resolved = resolve(bp, RenderContext(), mode=ResolveMode.KEYS)

    assert runs_of(resolved)[-1].text == placeholder_syntax("org_name")


def test_keys_mode_still_records_provenance(make_docx, fixture_provider):
    bp = blueprint_of(make_docx(paragraphs=["Hi {{ org_name }}"]), fixture_provider)

    runs = runs_of(resolve(bp, mode=ResolveMode.KEYS))

    assert runs[-1].placeholder_key == "org_name"


def test_keys_mode_expands_a_joined_placeholder_into_its_parts(make_docx, fixture_provider):
    bp = blueprint_of(
        make_docx(paragraphs=["{{ org_name, invoice_no, sep=', ' }}"]), fixture_provider,
    )

    runs = runs_of(resolve(bp, mode=ResolveMode.KEYS))

    assert [r.text for r in runs] == ["{{ org_name }}", ", ", "{{ invoice_no }}"]


def test_values_mode_is_unchanged_by_default(make_docx, fixture_provider):
    """The mode argument defaults to VALUES; existing callers keep their behaviour."""
    bp = blueprint_of(make_docx(paragraphs=["Hi {{ org_name }}"]), fixture_provider)

    resolved = DocumentResolver(
        values("org_name"), NoAssets(), DiagnosticCollector(),
    ).resolve(bp)

    assert runs_of(resolved)[-1].text == "<org_name>"


# --- render_raw --------------------------------------------------------------

def test_render_raw_produces_a_document_with_placeholders_intact(make_docx,
                                                                 fixture_provider, tmp_path):
    bp = blueprint_of(
        make_docx(paragraphs=["Invoice for {{ org_name }}", "No. {{ invoice_no }}"]),
        fixture_provider,
    )

    render = TemplateRenderingPipeline(NoAssets()).render_raw(bp)

    assert render.docx is not None
    out = tmp_path / "raw.docx"
    out.write_bytes(render.docx)

    text = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "{{ org_name }}" in text
    assert "{{ invoice_no }}" in text


def test_render_raw_reproduces_an_author_designed_row_template(make_docx, tmp_path):
    """A row-placeholder table round-trips: one row of column keys, as authored.

    The provider must register invl_* as COLUMN, otherwise the row is never
    promoted to a RowPlaceholder and this asserts nothing.
    """
    provider = column_provider()
    path = make_docx(table=[["#", "Item"], ["{{ invl_n }}", "{{ invl_desc }}"]])
    bp = blueprint_of(path, provider)

    assert row_placeholders_in(bp), "fixture did not produce a RowPlaceholder"

    render = TemplateRenderingPipeline(NoAssets()).render_raw(bp)

    out = tmp_path / "raw_table.docx"
    out.write_bytes(render.docx)
    table = Document(str(out)).tables[0]

    assert [c.text for c in table.rows[1].cells] == ["{{ invl_n }}", "{{ invl_desc }}"]


def test_a_row_placeholder_expands_once_in_keys_mode(make_docx):
    """Regression: KEYS mode had no table data, so author rows were skipped entirely."""
    bp = blueprint_of(
        make_docx(table=[["#", "Item"], ["{{ invl_n }}", "{{ invl_desc }}"]]),
        column_provider(),
    )

    resolved = resolve(bp, mode=ResolveMode.KEYS)
    table = [b for s in resolved.sections for b in s.blocks
             if not isinstance(b, ResolvedParagraph)][0]

    assert len(table.rows) == 2                      # header + one template row
    assert [c.blocks[0].runs[0].text for c in table.rows[1].cells] == [
        "{{ invl_n }}", "{{ invl_desc }}",
    ]


def test_a_row_placeholder_cell_keeps_its_provenance(make_docx):
    bp = blueprint_of(
        make_docx(table=[["#", "Item"], ["{{ invl_n }}", "{{ invl_desc }}"]]),
        column_provider(),
    )

    resolved = resolve(bp, mode=ResolveMode.KEYS)
    table = [b for s in resolved.sections for b in s.blocks
             if not isinstance(b, ResolvedParagraph)][0]

    assert [c.blocks[0].runs[0].placeholder_key for c in table.rows[1].cells] == [
        "invl_n", "invl_desc",
    ]


def test_keys_mode_needs_no_table_data_for_a_row_placeholder(make_docx):
    """No 'row_placeholder_no_data' warning, because KEYS mode supplies its own row."""
    bp = blueprint_of(
        make_docx(table=[["#", "Item"], ["{{ invl_n }}", "{{ invl_desc }}"]]),
        column_provider(),
    )
    diagnostics = DiagnosticCollector()

    DocumentResolver(
        RenderContext(), NoAssets(), diagnostics, ResolveMode.KEYS,
    ).resolve(bp)

    assert [d.code for d in diagnostics.items] == []


def test_render_raw_emits_no_diagnostics_for_a_complete_template(make_docx,
                                                                 fixture_provider):
    bp = blueprint_of(make_docx(paragraphs=["Invoice for {{ org_name }}"]),
                      fixture_provider)

    render = TemplateRenderingPipeline(NoAssets()).render_raw(bp)

    assert render.diagnostics.items == []
